"""Document parser routing PDF and DOCX into a unified ParsedDocument structure."""
import io
import re
import zipfile

from models.pipeline_models import ImageInfo, ParsedDocument, ParsedPage
from pipeline.pdf_parser import parse_pdf


def parse_document(file_bytes: bytes, filename: str | None) -> ParsedDocument:
    lower = (filename or "").lower()
    if file_bytes.startswith(b"%PDF"):
        return parse_pdf(file_bytes)
    if lower.endswith(".docx"):
        return parse_docx(file_bytes)
    if lower.endswith(".pdf") or not lower:
        return parse_pdf(file_bytes)
    return ParsedDocument(error=f"UNSUPPORTED_DOCUMENT_TYPE: {filename or 'unknown'}")


def parse_docx(docx_bytes: bytes) -> ParsedDocument:
    """Parse DOCX into virtual pages with paragraph-level locations.

    DOCX does not have real pages, so we assign virtual page numbers by grouping
    paragraphs. Images are associated with the paragraph that embeds them (or the
    nearest preceding paragraph) so that evidence location can show something more
    useful than "page 1".
    """
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as archive:
            if "word/document.xml" not in archive.namelist():
                return ParsedDocument(error="DOCX_PARSE_ERROR: missing word/document.xml")

            try:
                from docx import Document
                doc = Document(io.BytesIO(docx_bytes))
            except Exception:
                # Fallback to legacy regex parser if python-docx is unavailable.
                return _parse_docx_legacy(archive)

            # Extract text per paragraph, keeping empty paragraphs for accurate indexes.
            paragraphs = []
            for para_idx, para in enumerate(doc.paragraphs):
                raw_text = para.text or ""
                text = _normalize_docx_text(raw_text)
                paragraphs.append({
                    "index": para_idx,
                    "text": text,
                    "raw_text": raw_text,
                })

            # Map each embedded image to the paragraph that contains its reference.
            # Word stores drawing references as <w:drawing> or <w:pict> inside runs.
            image_paragraph_index = {}
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        image_bytes = rel.target_part.blob
                    except Exception:
                        continue
                    # Try to find which paragraph embeds this image by scanning part xml.
                    image_paragraph_index[rel.rId] = {
                        "bytes": image_bytes,
                        "paragraph_index": None,
                    }

            for para_idx, para in enumerate(doc.paragraphs):
                para_xml = para._p.xml
                for rId in image_paragraph_index:
                    if rId in para_xml:
                        image_paragraph_index[rId]["paragraph_index"] = para_idx

            # Collect images in document order and assign virtual pages.
            # Virtual page: every ~35 paragraphs or whenever an explicit page break appears.
            images_by_para = {}
            for rId, info in image_paragraph_index.items():
                para_idx = info["paragraph_index"]
                if para_idx is None:
                    para_idx = 0
                images_by_para.setdefault(para_idx, []).append(info["bytes"])

            # Build virtual pages.
            PAGE_BREAK_RE = re.compile(r"<w:br\s+[^>]*w:type=\"page\"|<w:br\s+[^>]*type=\"page\"")
            PARAGRAPHS_PER_VIRTUAL_PAGE = 35

            pages = []
            current_page_texts = []
            current_page_images = []
            current_page_num = 1
            current_para_count = 0

            for para_idx, para in enumerate(paragraphs):
                # Detect explicit page breaks in paragraph XML.
                para_xml = doc.paragraphs[para_idx]._p.xml
                has_page_break = bool(PAGE_BREAK_RE.search(para_xml))

                para_images = images_by_para.get(para_idx, [])
                if para_images:
                    for img_bytes in para_images:
                        current_page_images.append(ImageInfo(
                            page=current_page_num,
                            bbox=[],
                            image_bytes=img_bytes,
                            paragraph_index=para_idx,
                        ))

                current_page_texts.append(para["text"])
                current_para_count += 1

                if has_page_break or current_para_count >= PARAGRAPHS_PER_VIRTUAL_PAGE:
                    pages.append(ParsedPage(
                        page_num=current_page_num,
                        text="\n".join(current_page_texts).strip(),
                        images=current_page_images,
                    ))
                    current_page_texts = []
                    current_page_images = []
                    current_page_num += 1
                    current_para_count = 0

            if current_page_texts or current_page_images:
                pages.append(ParsedPage(
                    page_num=current_page_num,
                    text="\n".join(current_page_texts).strip(),
                    images=current_page_images,
                ))

            if not pages:
                pages.append(ParsedPage(page_num=1, text="", images=[]))

            return ParsedDocument(pages=pages)
    except Exception as e:
        return ParsedDocument(error=f"DOCX_PARSE_ERROR: {str(e)}")


def _parse_docx_legacy(archive: zipfile.ZipFile) -> ParsedDocument:
    """Fallback parser using regex, with no paragraph index information."""
    document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
    text_fragments = re.findall(r"<w:t[^>]*>(.*?)</w:t>", document_xml)
    text = _normalize_docx_text("".join(_decode_xml_entities(fragment) for fragment in text_fragments))

    images = []
    for name in archive.namelist():
        if not name.startswith("word/media/"):
            continue
        try:
            images.append(ImageInfo(page=1, bbox=[], image_bytes=archive.read(name)))
        except Exception:
            continue

    return ParsedDocument(pages=[ParsedPage(page_num=1, text=text, images=images)])


def _normalize_docx_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\u3000", " ")
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _decode_xml_entities(text: str) -> str:
    return (
        text.replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&amp;", "&")
        .replace("&quot;", '"')
        .replace("&apos;", "'")
    )
